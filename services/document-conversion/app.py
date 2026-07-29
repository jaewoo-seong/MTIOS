import base64
import csv
import hashlib
import io
import json
import os
import re
from typing import Any

import boto3
import fitz
import markdown as markdown_lib
import pdfplumber
import pytesseract
from bs4 import BeautifulSoup
from docx import Document
from fastapi import FastAPI, File, Form, Header, HTTPException, UploadFile
from fastapi.responses import Response
from PIL import Image
from pydantic import BaseModel
from weasyprint import HTML

app = FastAPI(title="MTI Document Conversion", version="1.0.0")
SERVICE_SECRET = os.getenv("DOCUMENT_CONVERSION_SERVICE_SECRET", "")
MAX_BYTES = 25 * 1024 * 1024


def authorize(value: str | None) -> None:
    if not SERVICE_SECRET:
        raise HTTPException(503, "DOCUMENT_CONVERSION_SERVICE_SECRET is not configured.")
    if value != f"Bearer {SERVICE_SECRET}":
        raise HTTPException(401, "Unauthorized.")


def storage_client():
    endpoint = os.getenv("RAILWAY_BUCKET_ENDPOINT")
    bucket = os.getenv("RAILWAY_BUCKET_NAME")
    access = os.getenv("RAILWAY_BUCKET_ACCESS_KEY_ID")
    secret = os.getenv("RAILWAY_BUCKET_SECRET_ACCESS_KEY")
    if not all([endpoint, bucket, access, secret]):
        return None, None
    client = boto3.client(
        "s3",
        endpoint_url=endpoint,
        aws_access_key_id=access,
        aws_secret_access_key=secret,
        region_name=os.getenv("RAILWAY_BUCKET_REGION", "auto"),
    )
    return client, bucket


def upload_asset(key: str, body: bytes, content_type: str) -> str | None:
    client, bucket = storage_client()
    if not client or not bucket:
        return None
    client.put_object(Bucket=bucket, Key=key, Body=body, ContentType=content_type)
    return key


@app.get("/health")
def health():
    return {
        "status": "ok",
        "service": "document-conversion",
        "ocr_languages": pytesseract.get_languages(config=""),
    }


@app.post("/v1/convert")
async def convert(
    file: UploadFile = File(...),
    languages: str = Form("en,ko"),
    authorization: str | None = Header(default=None),
):
    authorize(authorization)
    body = await file.read()
    if not body or len(body) > MAX_BYTES:
        raise HTTPException(413, "File is empty or exceeds 25MB.")
    name = file.filename or "document"
    mime = file.content_type or "application/octet-stream"
    extension = name.rsplit(".", 1)[-1].lower() if "." in name else ""
    if mime == "application/pdf" or extension == "pdf":
        return convert_pdf(name, body, languages)
    return convert_office_text(name, mime, body)


def convert_pdf(filename: str, body: bytes, languages: str):
    source_hash = hashlib.sha256(body).hexdigest()
    pdf = fitz.open(stream=body, filetype="pdf")
    plumber = pdfplumber.open(io.BytesIO(body))
    pages: list[dict[str, Any]] = []
    markdown_pages: list[str] = []
    warnings: list[str] = []
    used_ocr = False
    lang_codes = ["kor" if item.strip().lower().startswith("ko") else "eng"
                  for item in languages.split(",") if item.strip()]
    ocr_lang = "+".join(dict.fromkeys(lang_codes)) or "eng+kor"

    for index, page in enumerate(pdf):
        page_number = index + 1
        blocks = []
        digital = page.get_text("blocks", sort=True)
        digital_text = "\n".join(item[4].strip() for item in digital if item[4].strip())
        confidence = 96 if len(digital_text.strip()) >= 20 else 0
        extraction_method = "digital"

        if confidence == 0:
            used_ocr = True
            extraction_method = "ocr"
            pix = page.get_pixmap(matrix=fitz.Matrix(2.5, 2.5), alpha=False)
            image = Image.open(io.BytesIO(pix.tobytes("png")))
            data = pytesseract.image_to_data(
                image, lang=ocr_lang, output_type=pytesseract.Output.DICT
            )
            words = []
            scores = []
            for item_index, text in enumerate(data["text"]):
                text = text.strip()
                try:
                    score = float(data["conf"][item_index])
                except (TypeError, ValueError):
                    score = -1
                if not text or score < 0:
                    continue
                words.append(text)
                scores.append(score)
                blocks.append({
                    "type": "paragraph",
                    "position": len(blocks),
                    "text": text,
                    "bbox": {
                        "x": data["left"][item_index],
                        "y": data["top"][item_index],
                        "width": data["width"][item_index],
                        "height": data["height"][item_index],
                    },
                    "confidence": round(score),
                    "extractionMethod": "ocr",
                    "aiRepaired": False,
                })
            digital_text = " ".join(words)
            confidence = round(sum(scores) / len(scores)) if scores else 0
            if confidence < 70:
                warnings.append(f"Page {page_number} OCR confidence is {confidence}%.")
        else:
            for position, item in enumerate(digital):
                text = item[4].strip()
                if not text:
                    continue
                size = max((span["size"] for block in page.get_text("dict")["blocks"]
                            if "lines" in block for line in block["lines"]
                            for span in line["spans"]
                            if span["text"].strip() in text), default=10)
                blocks.append({
                    "type": "heading" if size >= 15 else "paragraph",
                    "position": position,
                    "text": text,
                    "bbox": {
                        "x": round(item[0]),
                        "y": round(item[1]),
                        "width": round(item[2] - item[0]),
                        "height": round(item[3] - item[1]),
                    },
                    "confidence": 96,
                    "extractionMethod": extraction_method,
                    "aiRepaired": False,
                })

        tables = []
        for table_position, table in enumerate(plumber.pages[index].find_tables()):
            cells = [[cell or "" for cell in row] for row in table.extract()]
            markdown_table = table_markdown(cells)
            tables.append({
                "position": table_position,
                "cells": cells,
                "bbox": {
                    "x": round(table.bbox[0]),
                    "y": round(table.bbox[1]),
                    "width": round(table.bbox[2] - table.bbox[0]),
                    "height": round(table.bbox[3] - table.bbox[1]),
                },
                "confidence": 85,
                "markdown": markdown_table,
            })

        preview = page.get_pixmap(matrix=fitz.Matrix(1.4, 1.4), alpha=False).tobytes("webp")
        preview_key = upload_asset(
            f"documents/previews/{source_hash}/page-{page_number}.webp",
            preview,
            "image/webp",
        )
        images = []
        for image_position, image_info in enumerate(page.get_images(full=True)):
            extracted = pdf.extract_image(image_info[0])
            image_bytes = extracted["image"]
            image_hash = hashlib.sha256(image_bytes).hexdigest()
            image_key = upload_asset(
                f"documents/images/{source_hash}/{image_hash}.{extracted['ext']}",
                image_bytes,
                f"image/{extracted['ext']}",
            )
            if image_key:
                images.append({
                    "position": image_position,
                    "bbox": None,
                    "storageKey": image_key,
                    "mimeType": f"image/{extracted['ext']}",
                    "width": extracted.get("width"),
                    "height": extracted.get("height"),
                    "altText": "",
                    "confidence": 100,
                })

        page_markdown = blocks_markdown(blocks)
        if tables:
            page_markdown += "\n\n" + "\n\n".join(item["markdown"] for item in tables)
        markdown_pages.append(f"## Page {page_number}\n\n{page_markdown}".strip())
        pages.append({
            "pageNumber": page_number,
            "width": round(page.rect.width),
            "height": round(page.rect.height),
            "text": digital_text,
            "confidence": confidence,
            "imageStorageKey": preview_key,
            "blocks": blocks,
            "tables": tables,
            "images": images,
        })

    plumber.close()
    markdown_text = "\n\n".join(markdown_pages)
    confidence = round(sum(page["confidence"] for page in pages) / len(pages)) if pages else 0
    return {
        "title": title_from_filename(filename),
        "markdown": markdown_text or "_No extractable text content was found._",
        "kind": "pdf",
        "pageCount": len(pages),
        "wordCount": len(markdown_text.split()),
        "truncated": False,
        "engine": "mti-layout-ocr",
        "engineVersion": "1.0.0",
        "language": detect_language(markdown_text),
        "ocrUsed": used_ocr,
        "confidence": confidence,
        "warnings": warnings,
        "needsReview": confidence < 80 or bool(warnings),
        "pages": pages,
    }


def convert_office_text(filename: str, mime: str, body: bytes):
    extension = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    tables = []
    if extension == "docx" or "wordprocessingml" in mime:
        document = Document(io.BytesIO(body))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            cells = [[cell.text for cell in row.cells] for row in table.rows]
            tables.append(cells)
        text = "\n\n".join(paragraphs)
        kind = "docx"
    elif extension in ("html", "htm") or "html" in mime:
        soup = BeautifulSoup(body.decode("utf-8", errors="replace"), "html.parser")
        for item in soup(["script", "style", "noscript"]):
            item.decompose()
        text = soup.get_text("\n", strip=True)
        kind = "html"
    elif extension in ("csv", "tsv"):
        rows = list(csv.reader(io.StringIO(body.decode("utf-8-sig")), delimiter="\t" if extension == "tsv" else ","))
        tables = [rows]
        text = table_markdown(rows)
        kind = "csv"
    elif extension == "json" or "json" in mime:
        text = "```json\n" + json.dumps(json.loads(body), ensure_ascii=False, indent=2) + "\n```"
        kind = "json"
    else:
        text = body.decode("utf-8", errors="replace")
        kind = "markdown" if extension in ("md", "markdown") else "text"
    markdown_text = text.strip()
    for item in tables:
        rendered = table_markdown(item)
        if rendered not in markdown_text:
            markdown_text += "\n\n" + rendered
    blocks = [{
        "type": "paragraph",
        "position": 0,
        "text": text,
        "bbox": None,
        "confidence": 95,
        "extractionMethod": "digital",
        "aiRepaired": False,
    }]
    return {
        "title": title_from_filename(filename),
        "markdown": markdown_text,
        "kind": kind,
        "pageCount": None,
        "wordCount": len(markdown_text.split()),
        "truncated": False,
        "engine": "mti-layout-ocr",
        "engineVersion": "1.0.0",
        "language": detect_language(markdown_text),
        "ocrUsed": False,
        "confidence": 95,
        "warnings": [],
        "needsReview": False,
        "pages": [{
            "pageNumber": 1,
            "width": None,
            "height": None,
            "text": text,
            "confidence": 95,
            "imageStorageKey": None,
            "blocks": blocks,
            "tables": [{
                "position": index,
                "cells": item,
                "bbox": None,
                "confidence": 90,
                "markdown": table_markdown(item),
            } for index, item in enumerate(tables)],
            "images": [],
        }],
    }


class ExportRequest(BaseModel):
    title: str
    markdown: str
    format: str


@app.post("/v1/export")
def export_document(payload: ExportRequest, authorization: str | None = Header(default=None)):
    authorize(authorization)
    if payload.format == "docx":
        document = Document()
        document.add_heading(payload.title, 0)
        for line in payload.markdown.splitlines():
            if line.startswith("# "):
                document.add_heading(line[2:], 1)
            elif line.startswith("## "):
                document.add_heading(line[3:], 2)
            elif line.strip():
                document.add_paragraph(re.sub(r"[*_`]", "", line))
        output = io.BytesIO()
        document.save(output)
        return Response(
            output.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    if payload.format == "pdf":
        body = markdown_lib.markdown(payload.markdown, extensions=["tables", "fenced_code"])
        html = f"""<!doctype html><html><head><meta charset="utf-8"><style>
        @page {{ size: A4; margin: 20mm; }}
        body {{ font-family: "Noto Sans CJK KR", "Noto Sans", sans-serif; line-height: 1.55; }}
        table {{ border-collapse: collapse; width: 100%; }}
        th, td {{ border: 1px solid #bbb; padding: 5px; }}
        </style></head><body><h1>{escape_html(payload.title)}</h1>{body}</body></html>"""
        return Response(HTML(string=html).write_pdf(), media_type="application/pdf")
    raise HTTPException(400, "Use pdf or docx.")


def table_markdown(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    padded = [row + [""] * (width - len(row)) for row in rows]
    header = "| " + " | ".join(escape_cell(item) for item in padded[0]) + " |"
    separator = "| " + " | ".join("---" for _ in range(width)) + " |"
    body = ["| " + " | ".join(escape_cell(item) for item in row) + " |" for row in padded[1:]]
    return "\n".join([header, separator, *body])


def blocks_markdown(blocks: list[dict[str, Any]]) -> str:
    output = []
    for block in blocks:
        text = block["text"].strip()
        if not text:
            continue
        output.append(f"### {text}" if block["type"] == "heading" else text)
    return "\n\n".join(output)


def title_from_filename(filename: str) -> str:
    value = re.sub(r"[_-]+", " ", filename.rsplit(".", 1)[0]).strip()
    return value[:1].upper() + value[1:] if value else "Untitled document"


def detect_language(value: str) -> str | None:
    korean = len(re.findall(r"[\uac00-\ud7af]", value))
    latin = len(re.findall(r"[A-Za-z]", value))
    if korean and latin:
        return "multilingual"
    if korean:
        return "ko"
    return "en" if latin else None


def escape_cell(value: str) -> str:
    return str(value).replace("|", "\\|").replace("\n", " ").strip()


def escape_html(value: str) -> str:
    return (
        value.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )
